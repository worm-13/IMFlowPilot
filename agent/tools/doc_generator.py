import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def set_font(run, font_name='宋体', font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

class DocGenerator:
    @staticmethod
    def generate_doc(topic: str, content: dict) -> str:
        doc = Document()

        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.add_run(topic)
        set_font(title_run, '黑体', 22, True)

        time_paragraph = doc.add_paragraph()
        time_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        time_run = time_paragraph.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        set_font(time_run, '宋体', 10)

        doc.add_page_break()

        for section, section_content in content.items():
            heading_paragraph = doc.add_heading(section, level=1)
            for run in heading_paragraph.runs:
                set_font(run, '黑体', 14, True)

            lines = section_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("- "):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                    for run in p.runs:
                        set_font(run, '宋体', 12)
                else:
                    p = doc.add_paragraph(line, style='Body Text')
                    for run in p.runs:
                        set_font(run, '宋体', 12)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        file_name = f"DOC_{timestamp}.docx"
        file_path = os.path.join(OUTPUT_DIR, file_name)
        doc.save(file_path)
        return file_path